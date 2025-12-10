# Docx_Translator
import os

from googletrans import Translator # to translate text
from pathlib import Path
from operator import index
from tkinter import Tk
from tkinter import filedialog as fd
from dotenv import load_dotenv # to load environment variables from .env file
from zipfile import BadZipFile # to handle invalid .docx files
import deepl # to translate text
from docx import Document   # to read and write .docx files


# Load the .env file from the same directory as this script
env_path = Path(__file__).parent / "Project_env.env"
load_dotenv(env_path) # load environment variables from .env file

# create google transalator obeject
translator = Translator()

# define a variable to select the file to be translated
def select_docx_file():
    """Open a file dialog and return the selected filepath to translate"""
    # create hidden root window
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)

    # create child window(file dialog)
    file_path = fd.askopenfilename( # assign a variable to open a dialog and select the file
        title="Choose a file to translate",
        filetypes=[
            ("Word Documents", "*.docx"), # shows only .docx files
            ("All FIles", "*.*") # show every type of file
        ],
        # use the environment variable to set the initial directory and a spare default value
        initialdir=os.getenv("docx_translator_dir", os.getenv("app_tools_dir")) 
    )
    
    # destroy the root dialog window
    root.destroy

    # Return None instead of empty string for better logic
    return file_path if file_path else None

def file_validation(file_path):
    """Validate if a file path was selected""" 
    if file_path is None: # when no file is selected
        return 
    elif not file_path.lower().endswith(".docx"): # validate file extension
        print("\nError!! The file selected must be a '.docx' file.")
        return 
    else:
        try: # validate file existence
            # When file is selected
            print(f"\nFile selected to translate: {file_path}", 
                    f"\nFile path: {os.path.dirname(file_path)}", 
                    f"\nFile name: {os.path.basename(file_path)}", 
                    f"\nFile size: {os.path.getsize(file_path)} KB", sep="")
            return True
        except FileExistsError: # file does not exist
            print(f"\nError!! The file {file_path} selected does not exist.")
            return 
        except BadZipFile as e: # file is not a valid .docx file
            print(f"\nError!! The file selected is not a valid .docx file: {e}")
            return
        except PermissionError as e: # file access permission error
            print(f"\nError!! Permission denied to access the file: {e}")
            return
        except Exception as e: # other errors
            print(f"\nError validating the file: {e}")
        return
        
def read_document(file_path):
    """Read the .docx file and return it as an object"""    
    selected_document = Document(file_path)
    doc = [] # create empty list to store paragraphs

    # extract all text paragraphs from the document
    try: 
        print("\nReading document content...")
        for paragraph in selected_document.paragraphs:
#            if paragraph.text.strip() != "": # skip empty paragraphs
            doc.append(paragraph.text)
        # print document content read success message
        print("\nDocument content read successfully.")
    except Exception as e: # handle errors while reading document
        print(f"Error reading the document: {e}")
        return

    # print paragraph count
    total = len(selected_document.paragraphs)
    print(f"Paragraph count: {total}\n")

    # print all paragraphs with a paragraph index as test
    for index, paragraph in enumerate(doc):
        if paragraph.strip() != "": # skip empty paragraphs
            print(index + 1, 
                  paragraph, f"Selected style: {selected_document.paragraphs[index].style.name}", 
                  f"Alignment: {selected_document.paragraphs[index].alignment}", 
                  f"Font: {selected_document.paragraphs[index].runs[0].font.name if selected_document.paragraphs[index].runs else "Default Font"}",
                  f"{len(selected_document.paragraphs[index].text)} characters",
                  sep = " - ")
#            print(f"P{index + 1}: {paragraph}") # alternative print format
        else:
            print(index + 1,"<Empty paragraph>", sep=" - ")
#            index - 1 # do not count empty paragraphs
    return selected_document if selected_document else None

def translate_text_googletrans(file_path, target_lang="ES"):
    """Translate text using googletrans module"""
    file_text = read_document(file_path)
    if file_text is None:
        print("The file selected does not exist or could not be read.")
        return

    translated_file = []
    print()
    for paragraph in file_text.paragraphs:
        if paragraph.text.strip() != "": # skip empty paragraphs
            translated = translator.translate(paragraph.text, 
                                              dest=target_lang)
            print(paragraph.text, " --> ", translated.text, sep="")
            translated_file.append(translated.text)
    print("\nThe file output is the following list:", 
          f"\n{translated_file}")
    print()
    return translated_file if translated_file else None

def transalted_doc_creation(file_path, translated_file):
    """Function to create an output file and 
    write translated text into the new file"""
    print("Creating output file...")
    
    # Create output file
    trans_file = Document()

    # Select original file path
    try:
        base_name = os.path.basename(file_path)
        name_no_ext, ext = os.path.splitext(base_name)
        print()
        print(f"Chosen base name --> {base_name}")
    except Exception as e: 
        print(f"Error reading the file: {e}")
        return
        

    # Create new file name
    new_name = name_no_ext + "_ES" + ext

    # Assign output file path
    try:
        output_dir = os.getenv("translated_docs_dir") or os.getenv("app_tools_dir") or os.getenv("trans_docs_dir") #or os.path.dirname(file_path)
        if not output_dir:
            print("Error!! No output directory defined in environment variables.")
            return
        output_path = os.path.join(output_dir, 
                                   new_name)
        print(f"Chosen file dir --> {output_dir}")
    except Exception as e:
        print(f"Error getting the output directory path: {e}")
        return



def main():
    """Main function to test file selection"""
    print("Select a .docx file to translate...")

    # get file selected from user
    chosen_file = select_docx_file()

    # Close if cancelled and no file selected
    if not chosen_file:
        print("\nUser closed the window before selecting a file.", 
              "\nGoodbye!")
        return

    # Validate file
    if file_validation(chosen_file) is None:
        return
    
    # Read document
    #selected_document = read_document(chosen_file)

    # Translate sample text
    #translated_text = translator.translate("Hello world", 
    #                                       dest='es').text
    #print(f"\nTranslated text: {translated_text}")

    # Translate document and save to translated files directory
    translated_file = translate_text_googletrans(chosen_file, target_lang="ES")

    # Call function to create a new document with the translated text
    transalted_doc_creation(chosen_file, translated_file)


if __name__=="__main__":
    main()